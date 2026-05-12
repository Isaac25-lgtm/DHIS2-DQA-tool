from __future__ import annotations

from datetime import UTC, datetime
from io import BytesIO
from uuid import UUID

import matplotlib
matplotlib.use("Agg")  # non-interactive backend
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Cm
from fastapi import HTTPException, status
from openpyxl import Workbook
from openpyxl.styles import Font
from sqlalchemy.orm import Session

from app.models.base import ExportStatus, ExportType, ReportStatus
from app.models.export_log import ExportLog
from app.models.report import Report
from app.models.user import User
from app.services.comment_sanitizer import sanitize_comment

try:  # pragma: no cover - optional dependency path
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
except Exception:  # pragma: no cover - optional dependency path
    canvas = None
    A4 = None

# ====================================================================
# Branding & colour constants
# ====================================================================
UCMB_PRIMARY_BLUE = RGBColor(15, 76, 129)       # 0F4C81
UCMB_ACCENT_GOLD = RGBColor(212, 175, 55)         # D4AF37
UCMB_LIGHT_BLUE_BG = "E8F2FA"
UCMB_WHITE = RGBColor(255, 255, 255)
UCMB_DARK_TEXT = RGBColor(30, 30, 30)

SCORE_GREEN = "27AE60"
SCORE_YELLOW = "D4A000"
SCORE_RED = "C82D2D"

SEVERITY_FILLS = {
    "CRITICAL": "E74C3C",
    "MAJOR": "E67E22",
    "MODERATE": "F4D03F",
    "MINOR": "3498DB",
    "MATCH": "27AE60",
    "EXACT_MATCH": "27AE60",
    "FLAGGED_ABOVE_5_PERCENT": "E67E22",
    "WITHIN_5_PERCENT": "3498DB",
    "NOT_AVAILABLE": "95A5A6",
    "MISSING": "E74C3C",
}
SEVERITY_DEFAULT_FILL = "95A5A6"
ROW_ALT_FILL = "F0F5FA"

# Matplotlib style colours
MPL_BLUE = "#0F4C81"
MPL_GOLD = "#D4AF37"
MPL_GREEN = "#27AE60"
MPL_RED = "#C82D2D"
MPL_ORANGE = "#E67E22"
MPL_YELLOW = "#F4D03F"
MPL_LIGHT_BLUE = "#3498DB"
MPL_GREY = "#95A5A6"


# ====================================================================
# Utility helpers
# ====================================================================

def _safe_filename(report: Report, export_type: ExportType) -> str:
    stem = report.title.replace(" ", "_").replace("/", "-")
    ext = export_type.value.lower()
    return f"{stem}.{ext}"


def _ensure_export_allowed(report: Report, export_type: ExportType) -> None:
    if export_type == ExportType.DOCX and report.status in {
        ReportStatus.GENERATED,
        ReportStatus.REVIEWED,
        ReportStatus.APPROVED,
        ReportStatus.EXPORTED,
    }:
        return
    if report.status not in {ReportStatus.APPROVED, ReportStatus.EXPORTED}:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Approve the report before exporting PDF or Excel. Word reports can be downloaded immediately after generation.",
        )


def _current_report_content(report: Report) -> str:
    return report.final_content or report.edited_content or report.generated_content


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()  # noqa: SLF001
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    tc_pr.append(shading)


def _set_cell_text_color(cell, color: RGBColor) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = color


def _style_table_header(row, fill: str = "0F4C81") -> None:
    for cell in row.cells:
        _set_cell_shading(cell, fill)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = UCMB_WHITE
                run.font.size = Pt(9)


def _coerce_number(value) -> float | None:
    if value is None:
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _format_value(value, *, as_percent: bool = False) -> str:
    if value is None:
        return "Not available"
    if isinstance(value, float):
        if value == int(value):
            text = str(int(value))
        else:
            text = f"{value:.1f}"
        return f"{text}%" if as_percent else text
    if as_percent and isinstance(value, int):
        return f"{value}%"
    return str(value)


def _label_implies_percent(label: str) -> bool:
    """Heuristic: a metric label like 'Completion percent' or 'Exact match rate' should
    render its numeric value with a trailing % sign."""
    if not label:
        return False
    lower = label.lower()
    return any(keyword in lower for keyword in ("percent", "rate", "%"))


def _rating_for(label: str, value: float | None) -> tuple[str, str]:
    """Return (rating_text, fill_color) for a numeric value given its label.

    Different label types use different scales. Counts (e.g. 'Critical discrepancy count',
    'Sync errors') are rated by 'lower is better' — 0 is green, anything > 0 is amber/red.
    Rates (e.g. 'Exact match rate') are rated by 'higher is better'."""
    if value is None:
        return ("Not assessed", "95A5A6")

    lower = label.lower()
    is_count = ("count" in lower or lower.endswith("errors") or lower.startswith("open ")
                or lower.startswith("overdue ") or "missing" in lower)
    is_lower_is_better = is_count or "discrepancy rate" in lower or "error" in lower

    if is_lower_is_better:
        if value <= 0:
            return ("Excellent", SCORE_GREEN)
        if value < 5:
            return ("Acceptable", SCORE_YELLOW)
        return ("Action needed", SCORE_RED)

    # Higher-is-better metrics (match rate, completion, etc.)
    if value >= 90:
        return ("Excellent", SCORE_GREEN)
    if value >= 70:
        return ("Good", SCORE_YELLOW)
    if value >= 50:
        return ("Needs improvement", SCORE_RED)
    return ("Poor", SCORE_RED)


def _score_color(percent: float | None) -> str:
    if percent is None:
        return "95A5A6"
    if percent >= 90:
        return SCORE_GREEN
    if percent >= 70:
        return SCORE_YELLOW
    return SCORE_RED


def _mpl_score_color(percent: float | None) -> str:
    return f"#{_score_color(percent)}"


def _severity_fill(severity: str | None) -> str:
    if not severity:
        return SEVERITY_DEFAULT_FILL
    return SEVERITY_FILLS.get(severity.upper().replace(" ", "_"), SEVERITY_DEFAULT_FILL)


def _add_paragraph_spacing(paragraph, before_pt: int = 0, after_pt: int = 4) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)


def _count_facilities(structured: dict) -> int:
    """Count how many facilities contributed data to this report."""
    rankings = structured.get("facility_score_ranking") or []
    if rankings:
        return len(rankings)
    # Fallback: count unique facilities from comparison_rows
    rows = structured.get("comparison_rows") or []
    names = {row.get("facility_name") for row in rows if row.get("facility_name")}
    if names:
        return len(names)
    # Facility-level report
    if structured.get("facility"):
        return 1
    coverage = structured.get("coverage") or {}
    return coverage.get("facilities_assessed", 0) or coverage.get("total_facilities_selected", 1)


# ====================================================================
# Chart helpers (matplotlib)
# ====================================================================

def _configure_chart_style() -> None:
    """Set global matplotlib style to match UCMB branding."""
    plt.rcParams.update({
        "font.family": "sans-serif",
        "font.sans-serif": ["Arial", "DejaVu Sans", "Helvetica"],
        "font.size": 9,
        "axes.titlesize": 11,
        "axes.titleweight": "bold",
        "axes.labelcolor": "#1E1E1E",
        "axes.edgecolor": "#CCCCCC",
        "axes.grid": True,
        "grid.alpha": 0.3,
        "grid.color": "#CCCCCC",
        "figure.facecolor": "white",
        "axes.facecolor": "white",
        "savefig.dpi": 150,
        "savefig.bbox": "tight",
        "savefig.pad_inches": 0.1,
    })


def _save_chart_to_stream(fig) -> BytesIO:
    stream = BytesIO()
    fig.savefig(stream, format="png", dpi=150, bbox_inches="tight", pad_inches=0.1)
    plt.close(fig)
    stream.seek(0)
    return stream


def _add_chart_to_document(document, chart_stream: BytesIO, width_inches: float = 5.8) -> None:
    """Insert a matplotlib chart as an inline image."""
    document.add_picture(chart_stream, width=Inches(width_inches))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_paragraph_spacing(last_paragraph, after_pt=10)


def _chart_facility_scores(structured: dict) -> BytesIO | None:
    """Horizontal bar chart: DQA scores per facility."""
    rankings = structured.get("facility_score_ranking") or []
    if not rankings:
        return None
    _configure_chart_style()
    names = [item.get("facility_name", "Unknown")[:25] for item in rankings]
    scores = [_coerce_number(item.get("dqa_score")) or 0 for item in rankings]
    categories = [item.get("score_category", "") for item in rankings]

    color_map = {"excellent": MPL_GREEN, "good": MPL_LIGHT_BLUE, "fair": MPL_YELLOW, "poor": MPL_ORANGE, "critical": MPL_RED, "needs improvement": MPL_ORANGE}
    bar_colors = [color_map.get((c or "").lower().strip(), MPL_GREY) for c in categories]

    fig, ax = plt.subplots(figsize=(8, max(2.5, len(names) * 0.4)))
    bars = ax.barh(range(len(names)), scores, color=bar_colors, edgecolor="white", height=0.7)
    ax.set_yticks(range(len(names)))
    ax.set_yticklabels(names, fontsize=8)
    ax.invert_yaxis()
    ax.set_xlabel("DQA Score (%)", fontsize=9)
    ax.set_title("DQA Scores by Facility", fontsize=12, color=MPL_BLUE, fontweight="bold")
    ax.set_xlim(0, 105)
    ax.axvline(x=90, color=MPL_GREEN, linestyle="--", alpha=0.5, linewidth=0.8, label="Excellent (90%)")
    ax.axvline(x=70, color=MPL_YELLOW, linestyle="--", alpha=0.5, linewidth=0.8, label="Fair (70%)")
    ax.legend(fontsize=7, loc="lower right")

    for bar, score in zip(bars, scores):
        ax.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height() / 2, f"{score:.0f}%", va="center", fontsize=8, fontweight="bold")

    fig.tight_layout()
    return _save_chart_to_stream(fig)


def _chart_severity_distribution(structured: dict) -> BytesIO | None:
    """Pie chart: severity distribution of discrepancies."""
    dqa_score = structured.get("dqa_score") or {}
    labels = ["Exact Match", "Minor", "Moderate", "Major", "Critical", "Missing"]
    keys = ["exact_count", "minor_count", "moderate_count", "major_count", "critical_count", "missing_count"]
    values = [_coerce_number(dqa_score.get(k)) or 0 for k in keys]
    total = sum(values)
    if total == 0:
        return None

    # Filter out zeros
    filtered = [(l, v) for l, v in zip(labels, values) if v > 0]
    if not filtered:
        return None
    labels, values = zip(*filtered)

    colors = [MPL_GREEN, MPL_LIGHT_BLUE, MPL_YELLOW, MPL_ORANGE, MPL_RED, MPL_GREY]
    color_subset = []
    for i, (label, _) in enumerate(zip(labels, values)):
        idx = ["Exact Match", "Minor", "Moderate", "Major", "Critical", "Missing"].index(label)
        color_subset.append(colors[idx])

    _configure_chart_style()
    fig, ax = plt.subplots(figsize=(6, 4.5))
    wedges, texts, autotexts = ax.pie(
        values, labels=None, autopct="%1.1f%%", startangle=140,
        colors=color_subset, pctdistance=0.75, explode=[0.05] * len(values),
    )
    for t in autotexts:
        t.set_fontsize(8)
        t.set_fontweight("bold")
    ax.set_title("Severity Distribution", fontsize=12, color=MPL_BLUE, fontweight="bold")
    ax.legend(wedges, labels, title="Severity", loc="center left", bbox_to_anchor=(1, 0.5), fontsize=8)
    fig.tight_layout()
    return _save_chart_to_stream(fig)


def _chart_completion_rates(structured: dict) -> BytesIO | None:
    """Bar chart: completion rates for the assessment scope."""
    coverage = structured.get("coverage") or {}
    districts = coverage.get("districts_covered") or []
    if not districts:
        return None

    _configure_chart_style()
    fig, ax = plt.subplots(figsize=(7, 3))
    # Simplified: show assessed vs pending
    assessed = coverage.get("facilities_assessed", 0) or 0
    total = coverage.get("total_facilities_selected", 0) or 0
    pending = coverage.get("facilities_pending", max(total - assessed, 0)) or 0
    pct = coverage.get("percentage_completed", 0) or 0

    categories = ["Completed", "Pending"]
    vals = [assessed, pending]
    bar_colors = [MPL_GREEN, MPL_GREY]
    bars = ax.bar(categories, vals, color=bar_colors, edgecolor="white", width=0.5)
    ax.set_ylabel("Number of Facilities", fontsize=9)
    ax.set_title(f"Assessment Completion: {_format_value(pct)}%", fontsize=12, color=MPL_BLUE, fontweight="bold")
    for bar, val in zip(bars, vals):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.2, str(int(val)), ha="center", fontweight="bold", fontsize=10)
    ax.set_ylim(0, max(vals) * 1.2 + 1)
    fig.tight_layout()
    return _save_chart_to_stream(fig)


def _chart_indicator_performance(structured: dict) -> BytesIO | None:
    """Horizontal bar chart of indicator exact match rates (top 15)."""
    findings = structured.get("indicator_findings") or []
    if not findings:
        return None
    findings_sorted = sorted(findings, key=lambda x: _coerce_number(x.get("exact_match_rate")) or 0)[:15]
    names = [f.get("indicator_name", "Unknown")[:35] for f in findings_sorted]
    rates = [_coerce_number(f.get("exact_match_rate")) or 0 for f in findings_sorted]

    _configure_chart_style()
    fig, ax = plt.subplots(figsize=(8, max(2.5, len(names) * 0.4)))
    bar_colors = [_mpl_score_color(r) for r in rates]
    bars = ax.barh(range(len(names)), rates, color=bar_colors, edgecolor="white", height=0.7)
    ax.set_yticks(range(len(names)))
    ax.set_yticklabels(names, fontsize=7)
    ax.invert_yaxis()
    ax.set_xlabel("Exact Match Rate (%)", fontsize=9)
    ax.set_title("Indicator Performance (Exact Match %)", fontsize=12, color=MPL_BLUE, fontweight="bold")
    ax.set_xlim(0, 105)
    ax.axvline(x=90, color=MPL_GREEN, linestyle="--", alpha=0.5, linewidth=0.8)
    for bar, rate in zip(bars, rates):
        ax.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height() / 2, f"{rate:.0f}%", va="center", fontsize=7, fontweight="bold")
    fig.tight_layout()
    return _save_chart_to_stream(fig)


def _chart_discrepancy_heatmap(structured: dict) -> BytesIO | None:
    """Facility-by-indicator heatmap using the maximum available percent difference."""
    rows = structured.get("comparison_rows") or []
    rows = [row for row in rows if row.get("facility_name") and row.get("indicator_name")]
    if not rows:
        return None

    facilities = sorted({row["facility_name"] for row in rows})
    indicators = sorted({f"{row.get('hmis_code', '')} {row['indicator_name']}".strip() for row in rows})
    if not facilities or not indicators:
        return None

    if len(facilities) > 25:
        facilities = facilities[:25]
    if len(indicators) > 20:
        indicators = indicators[:20]

    matrix = [[0.0 for _ in indicators] for _ in facilities]
    index_by_facility = {name: index for index, name in enumerate(facilities)}
    index_by_indicator = {name: index for index, name in enumerate(indicators)}
    for row in rows:
        indicator_label = f"{row.get('hmis_code', '')} {row['indicator_name']}".strip()
        if row["facility_name"] not in index_by_facility or indicator_label not in index_by_indicator:
            continue
        value = _coerce_number(row.get("max_percent_diff") or row.get("discrepancy_percent")) or 0.0
        matrix[index_by_facility[row["facility_name"]]][index_by_indicator[indicator_label]] = min(value, 50.0)

    _configure_chart_style()
    fig_width = max(7, len(indicators) * 0.45)
    fig_height = max(3.5, len(facilities) * 0.32)
    fig, ax = plt.subplots(figsize=(fig_width, fig_height))
    image = ax.imshow(matrix, cmap="RdYlGn_r", aspect="auto", vmin=0, vmax=20)
    ax.set_xticks(range(len(indicators)))
    ax.set_xticklabels([label[:22] for label in indicators], rotation=45, ha="right", fontsize=7)
    ax.set_yticks(range(len(facilities)))
    ax.set_yticklabels([label[:28] for label in facilities], fontsize=7)
    ax.set_title("Heat Map of Discrepancy Concentration", fontsize=12, color=MPL_BLUE, fontweight="bold")
    cbar = fig.colorbar(image, ax=ax, fraction=0.025, pad=0.02)
    cbar.set_label("Max percent difference", fontsize=8)
    fig.tight_layout()
    return _save_chart_to_stream(fig)


def _chart_discrepancy_types(structured: dict) -> BytesIO | None:
    """Grouped bar chart: Register-to-HMIS vs DHIS2 entry vs multiple-stage errors."""
    dist = structured.get("discrepancy_type_distribution") or {}
    if not dist:
        return None
    labels = ["Register→HMIS", "DHIS2 Entry", "Multi-Stage", "Missing"]
    keys = ["register_to_hmis_error_count", "dhis2_entry_error_count", "multiple_stage_error_count", "missing_value_count"]
    values = [_coerce_number(dist.get(k)) or 0 for k in keys]
    total = sum(values)
    if total == 0:
        return None

    _configure_chart_style()
    fig, ax = plt.subplots(figsize=(7, 3.5))
    bar_colors = [MPL_ORANGE, MPL_LIGHT_BLUE, MPL_RED, MPL_GREY]
    bars = ax.bar(labels, values, color=bar_colors, edgecolor="white", width=0.55)
    ax.set_ylabel("Count", fontsize=9)
    ax.set_title("Discrepancy Type Distribution", fontsize=12, color=MPL_BLUE, fontweight="bold")
    for bar, val in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.3, str(int(val)), ha="center", fontweight="bold", fontsize=9)
    ax.set_ylim(0, max(values) * 1.25 + 1)
    fig.tight_layout()
    return _save_chart_to_stream(fig)


# ====================================================================
# UCMB Logo (text-based crest)
# ====================================================================

def _add_logo_crest(document: Document) -> None:
    """Add a text-based UCMB logo/crest using a styled table."""
    crest_table = document.add_table(rows=1, cols=1)
    crest_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    crest_table.autofit = True
    cell = crest_table.rows[0].cells[0]
    _set_cell_shading(cell, "0F4C81")

    # Gold border effect via cell margins
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_paragraph_spacing(p, before_pt=8, after_pt=2)

    # UCMB acronym in gold
    run1 = p.add_run("  U  C  M  B  ")
    run1.font.size = Pt(18)
    run1.font.bold = True
    run1.font.color.rgb = UCMB_ACCENT_GOLD

    # Subtitle
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_paragraph_spacing(p2, before_pt=0, after_pt=8)
    run2 = p2.add_run("HMIS 105 Data Quality Assessment")
    run2.font.size = Pt(9)
    run2.font.color.rgb = UCMB_WHITE


# ====================================================================
# Page helpers (cover, header, footer)
# ====================================================================

def _add_cover_page(document: Document, report: Report, facility_count: int) -> None:
    """Professional cover page with UCMB branding."""
    # Some vertical space
    for _ in range(3):
        p = document.add_paragraph()
        _add_paragraph_spacing(p, before_pt=6, after_pt=4)

    # UCMB Logo/crest centered
    _add_logo_crest(document)
    document.add_paragraph()

    # Banner line
    banner = document.add_table(rows=1, cols=1)
    banner.autofit = True
    _set_cell_shading(banner.rows[0].cells[0], "0F4C81")
    banner_cell = banner.rows[0].cells[0]
    banner_p = banner_cell.paragraphs[0]
    banner_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    banner_run = banner_p.add_run("UGANDA CATHOLIC MEDICAL BUREAU")
    banner_run.font.size = Pt(13)
    banner_run.font.bold = True
    banner_run.font.color.rgb = UCMB_ACCENT_GOLD
    _add_paragraph_spacing(banner_p, before_pt=6, after_pt=6)

    document.add_paragraph()

    # Main title
    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(report.title)
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = UCMB_PRIMARY_BLUE
    _add_paragraph_spacing(title_p, before_pt=14, after_pt=2)

    # Subtitle
    sub_p = document.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("HMIS 105 Data Quality Assessment Report")
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = RGBColor(100, 110, 120)
    _add_paragraph_spacing(sub_p, before_pt=0, after_pt=20)

    # Horizontal rule
    rule = document.add_table(rows=1, cols=1)
    _set_cell_shading(rule.rows[0].cells[0], "0F4C81")
    rule.rows[0].cells[0].paragraphs[0].add_run(" ")
    for row in rule.rows:
        row.height = Cm(0.05)

    document.add_paragraph()

    # Metadata table
    structured = report.structured_input_json or {}
    org = structured.get("organization") or {}
    coverage = structured.get("coverage") or {}
    round_data = structured.get("assessment_round") or {}

    meta = document.add_table(rows=0, cols=2)
    meta.autofit = True
    metadata_pairs = [
        ("Report Type", report.report_type.value.replace("_", " ").title()),
        ("Date Generated", datetime.now(UTC).strftime("%d %B %Y")),
        ("Assessment Period", round_data.get("reporting_period", "N/A")),
        ("Prepared By", org.get("report_prepared_by", "UCMB DQA System")),
        ("Facilities Assessed", str(facility_count)),
        ("Status", "CONFIDENTIAL"),
    ]
    if report.assessment_round_id:
        metadata_pairs.insert(1, ("Assessment Round", round_data.get("name", str(report.assessment_round_id))))

    for label, value in metadata_pairs:
        row = meta.add_row()
        label_cell = row.cells[0]
        label_cell.text = label
        label_run = label_cell.paragraphs[0].runs[0]
        label_run.font.bold = True
        label_run.font.size = Pt(10)
        label_run.font.color.rgb = UCMB_PRIMARY_BLUE
        value_cell = row.cells[1]
        value_cell.text = str(value)
        value_run = value_cell.paragraphs[0].runs[0]
        value_run.font.size = Pt(10)

    document.add_paragraph()
    document.add_paragraph()

    # Table of contents placeholder
    toc_heading = document.add_paragraph()
    toc_run = toc_heading.add_run("TABLE OF CONTENTS")
    toc_run.font.bold = True
    toc_run.font.size = Pt(11)
    toc_run.font.color.rgb = UCMB_PRIMARY_BLUE
    _add_paragraph_spacing(toc_heading, before_pt=12, after_pt=6)

    toc_sections = [
        "1. Executive Snapshot",
        "2. Critical Chase List",
        "3. Scope and Method",
        "4. Main Findings",
        "5. Facility Performance",
        "6. Indicator Performance",
        "7. DHIS2 No-Data Review",
        "8. Source Document Review",
        "9. Corrective Action Plan",
        "10. Limitations and Appendix Notes",
    ]
    if facility_count > 1:
        toc_sections.insert(3, "3a. Facility Performance Details")
    for section in toc_sections:
        sp = document.add_paragraph(section)
        sp.runs[0].font.size = Pt(10)
        sp.runs[0].font.color.rgb = UCMB_DARK_TEXT
        _add_paragraph_spacing(sp, after_pt=2)

    # Confidential notice
    document.add_paragraph()
    conf_p = document.add_paragraph()
    conf_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_run = conf_p.add_run("CONFIDENTIAL – For Authorised Recipients Only")
    conf_run.font.size = Pt(9)
    conf_run.font.italic = True
    conf_run.font.color.rgb = RGBColor(180, 50, 50)

    document.add_page_break()


def _add_deepseek_header(document: Document) -> None:
    """Header with UCMB branding and page numbers."""
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.different_first_page_header_footer = True

    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""

    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ucmb_run = paragraph.add_run("UCMB  ")
    ucmb_run.font.bold = True
    ucmb_run.font.size = Pt(11)
    ucmb_run.font.color.rgb = UCMB_PRIMARY_BLUE
    sep_run = paragraph.add_run("|  ")
    sep_run.font.size = Pt(9)
    sep_run.font.color.rgb = RGBColor(160, 170, 180)
    dqa_run = paragraph.add_run("HMIS 105 DQA Report")
    dqa_run.font.size = Pt(9)
    dqa_run.font.color.rgb = UCMB_DARK_TEXT


def _add_footer(document: Document) -> None:
    """Footer with page numbers and confidentiality."""
    section = document.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    conf_run = footer.paragraphs[0].add_run("UCMB DQA Report  •  CONFIDENTIAL  •  Page ")
    conf_run.font.size = Pt(7)
    conf_run.font.color.rgb = RGBColor(140, 150, 160)
    conf_run.font.italic = True

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "

    run_page = footer.paragraphs[0].add_run()
    run_page.font.size = Pt(7)
    run_page.font.color.rgb = RGBColor(140, 150, 160)
    run_page._r.append(fld_char_begin)  # noqa: SLF001
    run_page._r.append(instr)           # noqa: SLF001
    run_page._r.append(fld_char_end)    # noqa: SLF001


# ====================================================================
# Metric summary boxes
# ====================================================================

def _add_metric_card(
    table,
    card_cell,
    label: str,
    value: str | float | None,
    score_percent: float | None = None,
) -> None:
    fill = _score_color(score_percent) if score_percent is not None else "E8F2FA"
    _set_cell_shading(card_cell, fill)

    p = card_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_paragraph_spacing(p, before_pt=6, after_pt=2)

    val_run = p.add_run(_format_value(value))
    val_run.font.size = Pt(22)
    val_run.font.bold = True
    val_run.font.color.rgb = UCMB_WHITE if score_percent is not None and score_percent >= 90 else UCMB_DARK_TEXT

    label_p = card_cell.add_paragraph()
    label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_paragraph_spacing(label_p, before_pt=0, after_pt=6)
    label_run = label_p.add_run(label)
    label_run.font.size = Pt(8)
    label_run.font.color.rgb = UCMB_WHITE if score_percent is not None and score_percent >= 90 else RGBColor(80, 90, 100)


def _add_metric_boxes(document: Document, report: Report) -> None:
    """Row of coloured summary metric cards."""
    structured = report.structured_input_json or {}

    document.add_heading("Executive Snapshot Metrics", level=1)

    cards: list[tuple[str, object, float | None]] = []

    score = structured.get("dqa_score") or {}
    if score.get("score_percent") is not None:
        cards.append(("Overall DQA Score", f"{_format_value(score['score_percent'])}%", _coerce_number(score.get("score_percent"))))

    summary = structured.get("summary") or {}
    if summary.get("exact_match_rate") is not None:
        cards.append(("Exact Match Rate", f"{_format_value(summary['exact_match_rate'])}%", _coerce_number(summary.get("exact_match_rate"))))

    coverage = structured.get("coverage") or {}
    if coverage.get("percentage_completed") is not None:
        cards.append(("Completion", f"{_format_value(coverage['percentage_completed'])}%", _coerce_number(coverage.get("percentage_completed"))))

    crit_count = None
    if score.get("critical_count") is not None:
        crit_count = score["critical_count"]
    elif summary.get("critical_discrepancy_count") is not None:
        crit_count = summary["critical_discrepancy_count"]
    if crit_count is not None:
        cards.append(("Critical Flags", crit_count, None))

    facility_count = _count_facilities(structured)
    if facility_count > 1:
        cards.append(("Facilities", facility_count, None))

    if not cards:
        return

    # Brief summary text
    intro_p = document.add_paragraph()
    _add_paragraph_spacing(intro_p, before_pt=4, after_pt=10)
    intro_run = intro_p.add_run(
        f"This report presents findings from the HMIS 105 Data Quality Assessment "
        f"covering {facility_count} health {'facilities' if facility_count > 1 else 'facility'}. "
        f"The assessment compares the source register count used as the primary verification reference, HMIS 105 reports, and DHIS2 entries "
        f"to identify discrepancies and recommend corrective actions."
    )
    intro_run.font.size = Pt(10)
    intro_run.font.color.rgb = UCMB_DARK_TEXT

    card_table = document.add_table(rows=1, cols=len(cards))
    card_table.autofit = True
    for idx, (label, value, score_pct) in enumerate(cards):
        cell = card_table.rows[0].cells[idx]
        _add_metric_card(card_table, cell, label, value, score_pct)

    document.add_paragraph()


# ====================================================================
# Tables with visual enhancements
# ====================================================================

def _add_key_value_table(document: Document, title: str, rows: list[tuple[str, object]]) -> None:
    if not rows:
        return
    document.add_heading(title, level=2)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Item"
    table.rows[0].cells[1].text = "Value"
    _style_table_header(table.rows[0])
    for row_idx, (label, value) in enumerate(rows):
        cells = table.add_row().cells
        cells[0].text = label
        cells[0].paragraphs[0].runs[0].font.bold = True
        cells[1].text = _format_value(value, as_percent=_label_implies_percent(label))
        if row_idx % 2 == 1:
            _set_cell_shading(cells[0], ROW_ALT_FILL)
            _set_cell_shading(cells[1], ROW_ALT_FILL)


def _add_score_card_table(document: Document, title: str, rows: list[tuple[str, float | int | None]], *, max_value: float = 100) -> None:
    chart_rows = [(label, _coerce_number(value)) for label, value in rows]
    chart_rows = [(label, value) for label, value in chart_rows if value is not None]
    if not chart_rows:
        return
    document.add_heading(title, level=2)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Measure"
    table.rows[0].cells[1].text = "Value"
    table.rows[0].cells[2].text = "Rating"
    _style_table_header(table.rows[0])

    for row_idx, (label, value) in enumerate(chart_rows):
        cells = table.add_row().cells
        cells[0].text = label
        cells[0].paragraphs[0].runs[0].font.bold = True
        cells[1].text = _format_value(value, as_percent=_label_implies_percent(label))

        rating_text, rating_fill = _rating_for(label, value)
        cells[2].text = rating_text
        _set_cell_shading(cells[2], rating_fill)
        if rating_fill in (SCORE_GREEN, SCORE_RED):
            for paragraph in cells[2].paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = UCMB_WHITE
                    run.font.bold = True

        if row_idx % 2 == 1:
            _set_cell_shading(cells[0], ROW_ALT_FILL)
            _set_cell_shading(cells[1], ROW_ALT_FILL)


def _add_data_table(document: Document, title: str, headers: list[str], rows: list[list[object]], *, limit: int = 30) -> None:
    if not rows:
        return
    document.add_heading(title, level=2)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    _style_table_header(table.rows[0])

    severity_col: int | None = None
    category_col: int | None = None
    for idx, h in enumerate(headers):
        hl = h.lower().strip()
        if hl in ("severity", "flag", "category"):
            severity_col = idx
        if hl in ("category", "score category"):
            category_col = idx

    for row_idx, data_row in enumerate(rows[:limit]):
        cells = table.add_row().cells
        for index, value in enumerate(data_row):
            cells[index].text = _format_value(value)
            if index == severity_col and isinstance(value, str):
                fill = _severity_fill(value)
                _set_cell_shading(cells[index], fill)
                if fill in (SEVERITY_FILLS.get("CRITICAL", "E74C3C"), SEVERITY_FILLS.get("MAJOR", "E67E22")):
                    _set_cell_text_color(cells[index], UCMB_WHITE)
                    for paragraph in cells[index].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
            elif index == category_col and isinstance(value, str):
                cat_fill = _score_color(
                    {"excellent": 95, "good": 85, "fair": 75, "poor": 50, "critical": 20}
                    .get(value.lower().strip(), None)
                )
                if value.lower().strip() in ("excellent", "good", "fair", "poor", "critical"):
                    _set_cell_shading(cells[index], cat_fill)
                    if cat_fill in (SCORE_GREEN, SCORE_RED):
                        _set_cell_text_color(cells[index], UCMB_WHITE)

        if row_idx % 2 == 1:
            for idx, cell in enumerate(cells):
                if idx != severity_col and idx != category_col:
                    existing = cell._tc.get_or_add_tcPr().findall(qn("w:shd"))  # noqa: SLF001
                    if not existing:
                        _set_cell_shading(cell, ROW_ALT_FILL)

    if len(rows) > limit:
        p = document.add_paragraph(
            f"Table truncated to first {limit} rows in Word export. Full data available in the system and XLSX export."
        )
        p.runs[0].font.italic = True
        p.runs[0].font.size = Pt(8)


# ====================================================================
# NEW: Field Notes & Context Section
# ====================================================================

def _build_field_notes_section(document: Document, structured: dict, facility_count: int) -> None:
    """Extract and present additional information / field notes from the assessment data.

    Every comment is run through the sanitizer before being placed in the report:
    insulting / profane comments are dropped entirely, and ALL-CAPS shouting is
    converted to readable sentence case. The report must remain audit-ready and safe
    to share with leadership and external stakeholders.
    """
    comments = []

    def _add_if_clean(facility: str, type_label: str, raw: str | None) -> None:
        cleaned = sanitize_comment(raw)
        if cleaned:
            comments.append({"facility": facility, "type": type_label, "comment": cleaned})

    # General facility comments (assessor's notes on each facility)
    for gc in structured.get("general_facility_comments") or []:
        _add_if_clean(
            gc.get("facility_name", "Unknown Facility"),
            "General Assessment Note",
            gc.get("comment"),
        )

    # Manager comments on assessment facilities
    for mc in structured.get("manager_comments") or []:
        _add_if_clean(
            mc.get("facility_name", "Unknown Facility"),
            "Manager Review Note",
            mc.get("comment"),
        )

    # Assessor comments from comparison rows
    for row in structured.get("comparison_rows") or []:
        _add_if_clean(
            row.get("facility_name", "Unknown Facility"),
            f"Field Note – {row.get('indicator_name', 'Indicator')}",
            row.get("assessor_comment"),
        )

    # Source document check comments
    for doc in structured.get("source_document_checks") or []:
        _add_if_clean(
            structured.get("facility", {}).get("facility_name", "Facility"),
            f"Source Document – {doc.get('source_document_name', 'Unknown')}",
            doc.get("comment"),
        )
    
    if not comments:
        # Generate context from structured data even without explicit notes
        coverage = structured.get("coverage") or {}
        round_data = structured.get("assessment_round") or {}
        facility_data = structured.get("facility") or {}
        
        context_lines = []
        if round_data.get("reporting_period"):
            context_lines.append(f"Assessment conducted for reporting period: {round_data['reporting_period']}.")
        if facility_data.get("administrative_area") or facility_data.get("district"):
            context_lines.append(f"Facility administrative area: {facility_data.get('administrative_area') or facility_data['district']}.")
        if facility_data.get("facility_type"):
            context_lines.append(f"Facility type: {facility_data['facility_type']}.")
        if coverage.get("administrative_areas_covered") or coverage.get("districts_covered"):
            areas = coverage.get("administrative_areas_covered") or coverage["districts_covered"]
            context_lines.append(f"Administrative areas covered: {', '.join(areas)}.")
        if coverage.get("facility_types"):
            context_lines.append(f"Facility types assessed: {', '.join(coverage['facility_types'])}.")
        if facility_data.get("ownership"):
            context_lines.append(f"Ownership: {facility_data['ownership']}.")
        
        if context_lines:
            document.add_heading("Field Notes & Assessment Context", level=1)
            intro_p = document.add_paragraph(
                "The following context was derived from the assessment data submitted by field assessors."
            )
            intro_p.runs[0].font.italic = True
            intro_p.runs[0].font.size = Pt(9)
            for line in context_lines:
                bp = document.add_paragraph(f"• {line}", style="List Bullet")
                bp.runs[0].font.size = Pt(9)
        return

    document.add_heading("Field Notes & Assessment Context", level=1)

    intro_p = document.add_paragraph(
        "The following notes and observations were recorded by field assessors during the "
        "data quality assessment. These provide important context for interpreting the findings."
    )
    intro_p.runs[0].font.italic = True
    intro_p.runs[0].font.size = Pt(9)
    _add_paragraph_spacing(intro_p, after_pt=10)

    # Group comments by facility
    facilities_order = []
    seen_facilities = set()
    for c in comments:
        if c["facility"] not in seen_facilities:
            seen_facilities.add(c["facility"])
            facilities_order.append(c["facility"])

    for facility in facilities_order:
        facility_comments = [c for c in comments if c["facility"] == facility]
        if facility_count > 1:
            document.add_heading(f"Notes: {facility}", level=2)
        else:
            document.add_heading("Assessment Observations", level=2)

        for fc in facility_comments:
            type_p = document.add_paragraph()
            type_run = type_p.add_run(fc["type"])
            type_run.font.bold = True
            type_run.font.size = Pt(9)
            type_run.font.color.rgb = UCMB_PRIMARY_BLUE
            _add_paragraph_spacing(type_p, before_pt=4, after_pt=2)

            comment_p = document.add_paragraph(fc["comment"])
            comment_p.runs[0].font.size = Pt(9)
            _add_paragraph_spacing(comment_p, after_pt=6)


# ====================================================================
# NEW: Key Findings Section
# ====================================================================

def _build_key_findings_section(document: Document, structured: dict, facility_count: int) -> None:
    """Generate structured key findings from the data."""
    document.add_heading("Key Findings", level=1)

    findings_list: list[tuple[str, str]] = []  # (severity_label, description)

    score = structured.get("dqa_score") or {}
    summary = structured.get("summary") or {}

    # Overall score finding
    overall_pct = _coerce_number(score.get("score_percent") or summary.get("exact_match_rate"))
    if overall_pct is not None:
        if overall_pct >= 90:
            findings_list.append(("EXCELLENT", f"Overall DQA score of {_format_value(overall_pct)}% exceeds the 90% excellent threshold."))
        elif overall_pct >= 70:
            findings_list.append(("SATISFACTORY", f"Overall DQA score of {_format_value(overall_pct)}% falls within the satisfactory range (70–89%)."))
        else:
            findings_list.append(("CONCERN", f"Overall DQA score of {_format_value(overall_pct)}% is below the 70% threshold and requires urgent attention."))

    # Critical discrepancies
    critical_count = _coerce_number(score.get("critical_count") or summary.get("critical_discrepancy_count"))
    if critical_count is not None and critical_count > 0:
        findings_list.append(("CRITICAL", f"{int(critical_count)} critical discrepancies identified – these represent the most severe data quality issues and require immediate corrective action."))
    elif critical_count == 0:
        findings_list.append(("POSITIVE", "No critical discrepancies were identified in this assessment."))

    # Major discrepancies
    major_count = _coerce_number(score.get("major_count") or 0)
    if major_count > 0:
        findings_list.append(("MAJOR", f"{int(major_count)} major discrepancies detected between reported and verified values."))

    # Exact match rate
    exact_rate = _coerce_number(summary.get("exact_match_rate"))
    if exact_rate is not None:
        if exact_rate >= 90:
            findings_list.append(("POSITIVE", f"Exact match rate of {_format_value(exact_rate)}% indicates strong alignment between data sources."))
        else:
            findings_list.append(("FLAG", f"Exact match rate of {_format_value(exact_rate)}% indicates significant data quality gaps."))

    # Source documents
    doc_rate = _coerce_number(summary.get("source_document_completeness_rate"))
    if doc_rate is not None:
        if doc_rate >= 90:
            findings_list.append(("POSITIVE", f"Source document completeness rate of {_format_value(doc_rate)}% shows good record-keeping practices."))
        else:
            findings_list.append(("FLAG", f"Source document completeness rate of {_format_value(doc_rate)}% indicates gaps in record availability or legibility."))

    # DHIS2 sync
    dhis2 = structured.get("dhis2_sync_summary") or {}
    sync_errors = _coerce_number(dhis2.get("dhis2_error_count")) or 0
    if sync_errors > 0:
        findings_list.append(("FLAG", f"{int(sync_errors)} DHIS2 synchronization errors detected, which may affect data completeness."))

    # Missing values
    missing = _coerce_number(score.get("missing_count") or 0)
    if missing > 0:
        findings_list.append(("FLAG", f"{int(missing)} values were missing across the indicators assessed."))

    # Open corrective actions
    open_actions = _coerce_number(summary.get("open_corrective_actions")) or 0
    overdue = _coerce_number(summary.get("overdue_corrective_actions")) or 0
    if open_actions > 0:
        findings_list.append(("INFO", f"{int(open_actions)} corrective actions remain open" + (f" ({int(overdue)} overdue)" if overdue > 0 else "") + "."))

    # Facility ranking findings (for multi-facility reports)
    facility_ranking = structured.get("facility_score_ranking") or []
    if len(facility_ranking) >= 2:
        best = facility_ranking[0]
        worst = facility_ranking[-1]
        findings_list.append(("INFO", f"Best performing facility: {best.get('facility_name', 'N/A')} ({_format_value(best.get('dqa_score'))}%)."))
        findings_list.append(("INFO", f"Lowest performing facility: {worst.get('facility_name', 'N/A')} ({_format_value(worst.get('dqa_score'))}%) – may require targeted support."))

    # Indicator findings
    indicator_findings = structured.get("indicator_findings") or []
    critical_indicators = [f for f in indicator_findings if _coerce_number(f.get("critical_discrepancy_count") or 0) > 0]
    if critical_indicators:
        names = [f.get("indicator_name", "Unknown")[:40] for f in critical_indicators[:3]]
        findings_list.append(("CRITICAL", f"Indicators with critical issues: {', '.join(names)}."))

    # Present findings
    severity_order = {"CRITICAL": 0, "MAJOR": 1, "FLAG": 2, "CONCERN": 3, "SATISFACTORY": 4, "POSITIVE": 5, "EXCELLENT": 6, "INFO": 7}
    findings_list.sort(key=lambda x: severity_order.get(x[0], 99))

    for sev, desc in findings_list:
        sev_fill = _severity_fill(sev)
        p = document.add_paragraph()
        badge_run = p.add_run(f"  {sev}  ")
        badge_run.font.size = Pt(8)
        badge_run.font.bold = True
        badge_run.font.color.rgb = UCMB_WHITE
        # Create a small inline badge – use a run highlight
        desc_run = p.add_run(f"  {desc}")
        desc_run.font.size = Pt(10)
        desc_run.font.color.rgb = UCMB_DARK_TEXT
        _add_paragraph_spacing(p, after_pt=6)

        # Simulate badge with background color on the badge text via a small table
        # Replace the paragraph approach with inline badge table
        p._element.getparent().remove(p._element)
        
        badge_table = document.add_table(rows=1, cols=2)
        badge_table.autofit = True
        # Remove table borders by setting no style
        badge_cell = badge_table.rows[0].cells[0]
        badge_cell.width = Inches(1.0)
        _set_cell_shading(badge_cell, _severity_fill(sev))
        badge_paragraph = badge_cell.paragraphs[0]
        badge_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        badge_run2 = badge_paragraph.add_run(sev)
        badge_run2.font.size = Pt(8)
        badge_run2.font.bold = True
        badge_run2.font.color.rgb = UCMB_WHITE
        
        desc_cell = badge_table.rows[0].cells[1]
        desc_cell.width = Inches(5.5)
        desc_paragraph = desc_cell.paragraphs[0]
        desc_run2 = desc_paragraph.add_run(desc)
        desc_run2.font.size = Pt(10)
        desc_run2.font.color.rgb = UCMB_DARK_TEXT
        
        # Add spacing after the table
        spacer = document.add_paragraph()
        _add_paragraph_spacing(spacer, before_pt=0, after_pt=2)


# ====================================================================
# NEW: Recommendations & Corrective Actions Section
# ====================================================================

def _build_recommendations_section(document: Document, structured: dict, facility_count: int) -> None:
    """Generate recommendations based on the data and present corrective actions."""
    document.add_heading("Recommendations & Corrective Actions", level=1)

    # Auto-generated recommendations based on data
    score = structured.get("dqa_score") or {}
    summary = structured.get("summary") or {}

    recommendations: list[tuple[str, str]] = []  # (priority, recommendation)

    overall_pct = _coerce_number(score.get("score_percent") or summary.get("exact_match_rate"))
    if overall_pct is not None and overall_pct < 90:
        if overall_pct < 70:
            recommendations.append(("URGENT", "Convene an emergency data quality review meeting with all facility data clerks and HMIS focal persons to address systemic data quality failures."))
        recommendations.append(("HIGH", "Implement a targeted training program for health facility staff on accurate HMIS data entry and reconciliation procedures."))
    
    critical_count = _coerce_number(score.get("critical_count") or summary.get("critical_discrepancy_count"))
    if critical_count is not None and critical_count > 0:
        recommendations.append(("HIGH", f"Immediately investigate and resolve the {int(critical_count)} critical discrepancies. Each critical finding should be assigned a responsible officer with a clear deadline for resolution."))

    # Source document quality
    doc_rate = _coerce_number(summary.get("source_document_completeness_rate"))
    if doc_rate is not None and doc_rate < 90:
        recommendations.append(("MEDIUM", "Improve source document management: ensure all registers are properly maintained, pages are numbered, and documents are stored securely for audit purposes."))

    # DHIS2 sync
    dhis2 = structured.get("dhis2_sync_summary") or {}
    sync_errors = _coerce_number(dhis2.get("dhis2_error_count")) or 0
    if sync_errors > 0:
        recommendations.append(("MEDIUM", f"Investigate and resolve DHIS2 synchronization errors ({int(sync_errors)} occurrences). Verify network connectivity and DHIS2 configuration at facility level."))

    # Multi-facility recommendations
    if facility_count > 1:
        facility_ranking = structured.get("facility_score_ranking") or []
        poor_facilities = [f for f in facility_ranking if (f.get("score_category") or "").lower() in ("poor", "critical", "needs improvement")]
        if poor_facilities:
            names = ", ".join(f.get("facility_name", "Unknown") for f in poor_facilities[:5])
            recommendations.append(("HIGH", f"Conduct focused supportive supervision visits to low-performing facilities: {names}."))

    # General recommendations
    recommendations.append(("MEDIUM", "Establish a regular (quarterly) DQA cycle to monitor progress on data quality improvements over time."))
    recommendations.append(("MEDIUM", "Integrate DQA findings into routine health management review meetings to ensure accountability."))
    recommendations.append(("LOW", "Document lessons learned from this assessment to inform future DQA rounds and refine the assessment methodology."))

    # Present recommendations
    priority_order = {"URGENT": 0, "HIGH": 1, "MEDIUM": 2, "LOW": 3}
    recommendations.sort(key=lambda x: priority_order.get(x[0], 99))

    rec_heading = document.add_heading("Management Recommendations", level=2)
    for priority, rec in recommendations:
        fill = {"URGENT": "E74C3C", "HIGH": "E67E22", "MEDIUM": "F4D03F", "LOW": "3498DB"}.get(priority, "95A5A6")
        p = document.add_paragraph()
        badge_run = p.add_run(f" {priority} ")
        badge_run.font.size = Pt(7)
        badge_run.font.bold = True
        p.add_run(f"  {rec}").font.size = Pt(10)
        _add_paragraph_spacing(p, after_pt=5)

    # Corrective actions table from structured data
    corrective_actions = structured.get("corrective_actions") or []
    if corrective_actions:
        document.add_heading("Corrective Actions Register", level=2)
        actions_table = document.add_table(rows=1, cols=5)
        actions_table.style = "Table Grid"
        for idx, header in enumerate(["Action", "Severity", "Status", "Responsible", "Facility"]):
            actions_table.rows[0].cells[idx].text = header
        _style_table_header(actions_table.rows[0])

        for row_idx, action in enumerate(corrective_actions):
            cells = actions_table.add_row().cells
            cells[0].text = str(action.get("action_description", ""))[:80]
            sev = str(action.get("severity", ""))
            cells[1].text = sev
            fill = _severity_fill(sev)
            _set_cell_shading(cells[1], fill)
            if fill in ("E74C3C", "E67E22"):
                _set_cell_text_color(cells[1], UCMB_WHITE)
            cells[2].text = str(action.get("status", "OPEN"))
            cells[3].text = str(action.get("responsible_person", "Not assigned"))
            cells[4].text = str(action.get("facility_name", "N/A"))
            if row_idx % 2 == 1:
                for idx in [0, 2, 3, 4]:
                    _set_cell_shading(cells[idx], ROW_ALT_FILL)

    # Source document audit for facility reports
    source_check_rows = structured.get("source_document_checks") or []
    if source_check_rows:
        document.add_heading("Source Document Audit", level=2)
        doc_headers = ["Document", "Available", "Complete", "Legible", "Missing Pages"]
        doc_table = document.add_table(rows=1, cols=len(doc_headers))
        doc_table.style = "Table Grid"
        for idx, h in enumerate(doc_headers):
            doc_table.rows[0].cells[idx].text = h
        _style_table_header(doc_table.rows[0])
        for idx, item in enumerate(source_check_rows):
            cells = doc_table.add_row().cells
            cells[0].text = str(item.get("source_document_name", ""))
            cells[1].text = "✓" if item.get("available") else "✗"
            cells[2].text = "✓" if item.get("complete") else "✗"
            cells[3].text = "✓" if item.get("legible") else "✗"
            cells[4].text = _format_value(item.get("missing_pages", 0) or 0)
            if item.get("comment"):
                comment_row = doc_table.add_row()
                comment_row.cells[0].text = "  ↳ Comment:"
                comment_cell = comment_row.cells[1]
                comment_cell_merged = comment_row.cells[1]
                # Write comment across remaining cells manually
                comment_cell_merged.text = str(item["comment"])[:120]
                for c_idx in range(2, 5):
                    comment_row.cells[c_idx].text = ""
            if idx % 2 == 1:
                for c in cells:
                    _set_cell_shading(c, ROW_ALT_FILL)


# ====================================================================
# Enhanced Statistical Dashboard (with charts)
# ====================================================================

# ====================================================================
# NEW: Blended narrative + tables + charts renderer
# ====================================================================

def _add_section_heading(document: Document, title: str) -> None:
    h = document.add_heading(title, level=1)
    for run in h.runs:
        run.font.color.rgb = UCMB_PRIMARY_BLUE


def _add_narrative_paragraphs(document: Document, narrative: str | None) -> None:
    if not narrative or not narrative.strip():
        return
    # Split on blank lines so multi-paragraph narratives render cleanly.
    paragraphs = [p.strip() for p in narrative.replace("\r\n", "\n").split("\n\n")]
    paragraphs = [p for p in paragraphs if p]
    for para_text in paragraphs:
        # The AI may include a newline inside a paragraph for line wrap; flatten that.
        flat = " ".join(line.strip() for line in para_text.split("\n") if line.strip())
        if not flat:
            continue
        p = document.add_paragraph(flat)
        for run in p.runs:
            run.font.size = Pt(11)
        _add_paragraph_spacing(p, before_pt=0, after_pt=8)


def _render_section_executive_summary(document, structured, sections, facility_count):
    _add_section_heading(document, "Executive Summary")
    _add_narrative_paragraphs(document, sections.get("executive_summary"))


def _render_section_scope(document, structured, sections, facility_count):
    _add_section_heading(document, "Assessment Scope and Coverage")
    _add_narrative_paragraphs(document, sections.get("scope_and_coverage"))
    coverage = structured.get("coverage") or {}
    if coverage:
        _add_key_value_table(
            document,
            "Coverage Summary",
            [
                ("Facilities selected", coverage.get("total_facilities_selected")),
                ("Facilities assessed", coverage.get("facilities_assessed")),
                ("Facilities pending", coverage.get("facilities_pending")),
                ("Completion percent", coverage.get("percentage_completed")),
                ("Administrative areas covered", ", ".join(coverage.get("administrative_areas_covered") or coverage.get("districts_covered") or [])),
                ("Facility types", ", ".join(coverage.get("facility_types") or [])),
            ],
        )
    chart = _chart_completion_rates(structured)
    if chart:
        _add_chart_to_document(document, chart)


def _render_section_methods(document, structured, sections, facility_count):
    _add_section_heading(document, "Methods and Data Sources")
    _add_narrative_paragraphs(document, sections.get("methods"))
    summary = structured.get("summary") or {}
    if summary:
        _add_score_card_table(
            document,
            "Round-Level Quality Indicators",
            [
                ("Exact match rate", summary.get("exact_match_rate")),
                ("Major discrepancy rate", summary.get("major_discrepancy_rate")),
                ("Source document completeness rate", summary.get("source_document_completeness_rate")),
                ("Critical discrepancy count", summary.get("critical_discrepancy_count")),
                ("Open corrective actions", summary.get("open_corrective_actions")),
                ("Overdue corrective actions", summary.get("overdue_corrective_actions")),
            ],
        )


def _render_section_overall_findings(document, structured, sections, facility_count):
    _add_section_heading(document, "Overall Statistical Findings")
    _add_narrative_paragraphs(document, sections.get("overall_findings"))
    score = structured.get("dqa_score") or {}
    if score:
        _add_score_card_table(
            document,
            "DQA Score and Severity Distribution",
            [
                ("DQA score percent", score.get("score_percent")),
                ("Exact matches", score.get("exact_count")),
                ("Minor discrepancies", score.get("minor_count")),
                ("Moderate discrepancies", score.get("moderate_count")),
                ("Major discrepancies", score.get("major_count")),
                ("Critical discrepancies", score.get("critical_count")),
                ("Missing values", score.get("missing_count")),
            ],
            max_value=max(_coerce_number(score.get("possible_points")) or 100, 100),
        )
    severity_chart = _chart_severity_distribution(structured)
    if severity_chart:
        _add_chart_to_document(document, severity_chart)
    discrepancy_chart = _chart_discrepancy_types(structured)
    if discrepancy_chart:
        _add_chart_to_document(document, discrepancy_chart)
    heatmap = _chart_discrepancy_heatmap(structured)
    if heatmap:
        h = document.add_heading("Heat Map of Discrepancy Concentration", level=2)
        for run in h.runs:
            run.font.color.rgb = UCMB_PRIMARY_BLUE
        _add_chart_to_document(document, heatmap, width_inches=6.5)


def _render_section_facility_performance(document, structured, sections, facility_count):
    _add_section_heading(document, "Facility Performance")
    _add_narrative_paragraphs(document, sections.get("facility_performance"))
    if facility_count > 1:
        chart = _chart_facility_scores(structured)
        if chart:
            _add_chart_to_document(document, chart)
    facility_rows = [
        [
            item.get("facility_name"),
            item.get("dqa_score"),
            item.get("score_category"),
            item.get("exact_count"),
            item.get("critical_count"),
            item.get("open_corrective_actions"),
        ]
        for item in structured.get("facility_score_ranking", [])
    ]
    _add_data_table(
        document,
        "Facility Performance Table",
        ["Facility", "DQA score", "Category", "Exact", "Critical", "Open actions"],
        facility_rows,
    )


def _render_section_indicator_findings(document, structured, sections, facility_count):
    _add_section_heading(document, "Indicator-Level Findings")
    _add_narrative_paragraphs(document, sections.get("indicator_findings"))
    chart = _chart_indicator_performance(structured)
    if chart:
        _add_chart_to_document(document, chart)
    indicator_rows = [
        [
            item.get("indicator_name"),
            item.get("hmis_code"),
            item.get("exact_match_rate"),
            item.get("major_discrepancy_count"),
            item.get("critical_discrepancy_count"),
            item.get("worst_facilities") if isinstance(item.get("worst_facilities"), str) else ", ".join(item.get("worst_facilities") or []),
        ]
        for item in structured.get("indicator_findings", [])
    ]
    _add_data_table(
        document,
        "Indicator Statistical Findings",
        ["Indicator", "HMIS code", "Exact rate", "Major", "Critical", "Worst facilities"],
        indicator_rows,
    )


def _render_section_dhis2_sync(document, structured, sections, facility_count):
    _add_section_heading(document, "DHIS2 Synchronization Findings")
    _add_narrative_paragraphs(document, sections.get("dhis2_synchronization"))
    dhis2 = structured.get("dhis2_sync_summary") or {}
    if dhis2:
        _add_score_card_table(
            document,
            "DHIS2 Synchronization Summary",
            [
                ("Values successfully pulled", dhis2.get("dhis2_values_successfully_pulled")),
                ("No-data responses", dhis2.get("dhis2_no_data_count")),
                ("Sync errors", dhis2.get("dhis2_error_count")),
            ],
            max_value=max(
                (_coerce_number(dhis2.get("dhis2_values_successfully_pulled")) or 0)
                + (_coerce_number(dhis2.get("dhis2_no_data_count")) or 0)
                + (_coerce_number(dhis2.get("dhis2_error_count")) or 0),
                1,
            ),
        )
        if dhis2.get("last_sync_time"):
            p = document.add_paragraph(f"Last DHIS2 sync time: {dhis2['last_sync_time']}")
            p.runs[0].font.size = Pt(8)
            p.runs[0].font.italic = True


def _render_section_source_documents(document, structured, sections, facility_count):
    _add_section_heading(document, "Source Document Findings")
    _add_narrative_paragraphs(document, sections.get("source_documents"))
    checks = structured.get("source_document_checks") or []
    if checks:
        rows = [
            [
                item.get("source_document_name"),
                "Yes" if item.get("available") else "No",
                "Yes" if item.get("complete") else "No",
                "Yes" if item.get("legible") else "No",
                "No" if item.get("missing_pages") else "Yes",
            ]
            for item in checks
        ]
        _add_data_table(
            document,
            "Source Document Checklist",
            ["Source Document", "Available", "Complete", "Legible", "All Pages Present"],
            rows,
        )


def _render_section_root_causes(document, structured, sections, facility_count):
    _add_section_heading(document, "Root-Cause Interpretation")
    _add_narrative_paragraphs(document, sections.get("root_causes"))
    comparison_rows = [
        [
            row.get("facility_name"),
            row.get("indicator_name"),
            row.get("hmis_code"),
            row.get("register_value"),
            row.get("hmis105_value"),
            row.get("dhis2_value_at_assessment") or row.get("dhis2_value"),
            row.get("severity"),
            row.get("issue_type"),
        ]
        for row in structured.get("comparison_rows", [])
    ]
    _add_data_table(
        document,
        "Detailed Comparison Rows",
        ["Facility", "Indicator", "HMIS", "Register", "HMIS 105", "DHIS2", "Severity", "Issue"],
        comparison_rows,
    )


def _render_section_comments_context(document, structured, sections, facility_count):
    _add_section_heading(document, "Team Comments and Contextual Observations")
    _add_narrative_paragraphs(document, sections.get("comments_context"))
    # Also surface the sanitized field-notes catalogue underneath the AI summary
    _build_field_notes_section(document, structured, facility_count)


def _render_section_corrective_actions(document, structured, sections, facility_count):
    _add_section_heading(document, "Corrective Action Plan")
    _add_narrative_paragraphs(document, sections.get("corrective_action_plan"))
    actions = structured.get("corrective_actions") or []
    if actions:
        rows = [
            [
                action.get("title") or action.get("description", "")[:60],
                action.get("severity") or action.get("priority"),
                action.get("status"),
                action.get("responsible_person"),
                action.get("due_date"),
            ]
            for action in actions
        ]
        _add_data_table(
            document,
            "Logged Corrective Actions",
            ["Action", "Severity", "Status", "Responsible", "Due"],
            rows,
        )


def _render_section_recommendations(document, structured, sections, facility_count):
    _add_section_heading(document, "Recommendations")
    _add_narrative_paragraphs(document, sections.get("recommendations"))
    # Also include the structured recommendations bullets the existing helper produces
    _build_recommendations_section(document, structured, facility_count)


def _render_section_limitations(document, structured, sections, facility_count):
    _add_section_heading(document, "Limitations")
    _add_narrative_paragraphs(document, sections.get("limitations"))


def _render_section_conclusion(document, structured, sections, facility_count):
    _add_section_heading(document, "Conclusion")
    _add_narrative_paragraphs(document, sections.get("conclusion"))


_BLENDED_SECTION_RENDERERS = [
    ("executive_summary", _render_section_executive_summary, False),
    ("scope_and_coverage", _render_section_scope, False),
    ("methods", _render_section_methods, False),
    ("overall_findings", _render_section_overall_findings, True),  # page break after
    ("facility_performance", _render_section_facility_performance, False),
    ("indicator_findings", _render_section_indicator_findings, True),
    ("dhis2_synchronization", _render_section_dhis2_sync, False),
    ("source_documents", _render_section_source_documents, False),
    ("root_causes", _render_section_root_causes, True),
    ("comments_context", _render_section_comments_context, False),
    ("corrective_action_plan", _render_section_corrective_actions, False),
    ("recommendations", _render_section_recommendations, False),
    ("limitations", _render_section_limitations, False),
    ("conclusion", _render_section_conclusion, False),
]


def _render_blended_report(document: Document, report: Report, facility_count: int) -> None:
    """Walk through every narrative section and interleave the AI's prose with the
    matching data tables and charts. This produces a report that reads like a
    coherent document instead of a dashboard followed by a separate essay."""
    structured = report.structured_input_json or {}
    sections = structured.get("narrative_sections") or {}

    document.add_page_break()
    for section_key, renderer, page_break_after in _BLENDED_SECTION_RENDERERS:
        # Skip sections whose data is empty AND whose narrative is empty
        if not sections.get(section_key) and not _section_has_data(structured, section_key):
            continue
        renderer(document, structured, sections, facility_count)
        if page_break_after:
            document.add_page_break()


def _section_has_data(structured: dict, section_key: str) -> bool:
    """Return True if a section has any underlying data worth rendering even when
    the AI narrative for that section is empty."""
    if section_key in {"executive_summary", "limitations", "conclusion", "root_causes",
                       "facility_performance", "indicator_findings", "overall_findings",
                       "scope_and_coverage", "methods"}:
        return True  # always render these scaffolding sections
    if section_key == "dhis2_synchronization":
        return bool(structured.get("dhis2_sync_summary"))
    if section_key == "source_documents":
        return bool(structured.get("source_document_checks"))
    if section_key == "comments_context":
        return any([
            structured.get("general_facility_comments"),
            structured.get("manager_comments"),
            any((row.get("assessor_comment") or row.get("manager_comment")) for row in structured.get("comparison_rows") or []),
        ])
    if section_key == "corrective_action_plan":
        return bool(structured.get("corrective_actions"))
    if section_key == "recommendations":
        return True
    return False


def _add_statistical_dashboard(document: Document, report: Report, facility_count: int) -> None:
    structured = report.structured_input_json or {}
    document.add_heading("Statistical Dashboard", level=1)

    # --- Coverage ---
    coverage = structured.get("coverage") or {}
    if coverage:
        _add_key_value_table(
            document,
            "Assessment Coverage",
            [
                ("Facilities selected", coverage.get("total_facilities_selected")),
                ("Facilities assessed", coverage.get("facilities_assessed")),
                ("Facilities pending", coverage.get("facilities_pending")),
                ("Completion percent", coverage.get("percentage_completed")),
                ("Administrative areas covered", ", ".join(coverage.get("administrative_areas_covered") or coverage.get("districts_covered") or [])),
                ("Facility types", ", ".join(coverage.get("facility_types") or [])),
            ],
        )

    # --- Completion chart ---
    completion_chart = _chart_completion_rates(structured)
    if completion_chart:
        _add_chart_to_document(document, completion_chart)

    # --- DQA Score ---
    if structured.get("dqa_score"):
        score = structured["dqa_score"]
        _add_score_card_table(
            document,
            "DQA Score and Severity Distribution",
            [
                ("DQA score percent", score.get("score_percent")),
                ("Exact matches", score.get("exact_count")),
                ("Minor discrepancies", score.get("minor_count")),
                ("Moderate discrepancies", score.get("moderate_count")),
                ("Major discrepancies", score.get("major_count")),
                ("Critical discrepancies", score.get("critical_count")),
                ("Missing values", score.get("missing_count")),
            ],
            max_value=max(_coerce_number(score.get("possible_points")) or 100, 100),
        )

    # --- Severity pie chart ---
    severity_chart = _chart_severity_distribution(structured)
    if severity_chart:
        _add_chart_to_document(document, severity_chart)

    # --- Round-level summary ---
    if structured.get("summary"):
        summary = structured["summary"]
        _add_score_card_table(
            document,
            "Round-Level Quality Indicators",
            [
                ("Exact match rate", summary.get("exact_match_rate")),
                ("Major discrepancy rate", summary.get("major_discrepancy_rate")),
                ("Source document completeness rate", summary.get("source_document_completeness_rate")),
                ("Critical discrepancy count", summary.get("critical_discrepancy_count")),
                ("Open corrective actions", summary.get("open_corrective_actions")),
                ("Overdue corrective actions", summary.get("overdue_corrective_actions")),
            ],
        )

    # --- Facility Scores Chart (multi-facility only) ---
    if facility_count > 1:
        facility_chart = _chart_facility_scores(structured)
        if facility_chart:
            _add_chart_to_document(document, facility_chart)

    # --- Discrepancy type distribution ---
    discrepancy_chart = _chart_discrepancy_types(structured)
    if discrepancy_chart:
        _add_chart_to_document(document, discrepancy_chart)

    heatmap_chart = _chart_discrepancy_heatmap(structured)
    if heatmap_chart:
        document.add_heading("Heat Map of Discrepancy Concentration", level=2)
        heatmap_note = document.add_paragraph(
            "The heat map highlights where discrepancies are concentrated across facilities and indicators. "
            "Darker red cells represent higher maximum percent differences and should be prioritised for reconciliation."
        )
        heatmap_note.runs[0].font.size = Pt(9)
        _add_chart_to_document(document, heatmap_chart, width_inches=6.5)

    # --- Comparison outcomes ---
    comparison_summary = structured.get("comparison_summary") or {}
    if comparison_summary:
        _add_score_card_table(
            document,
            "Comparison Outcome Summary",
            [
                ("Exact matches", comparison_summary.get("exact_matches")),
                ("Within 5 percent", comparison_summary.get("within_5_percent")),
                ("Flagged above 5 percent", comparison_summary.get("flagged_above_5_percent")),
                ("Critical flags", comparison_summary.get("critical_flags")),
                ("Incomplete rows", comparison_summary.get("incomplete_rows")),
            ],
            max_value=max(_coerce_number(comparison_summary.get("total_rows_assessed")) or 1, 1),
        )

    # --- DHIS2 sync ---
    dhis2 = structured.get("dhis2_sync_summary") or {}
    if dhis2:
        _add_score_card_table(
            document,
            "DHIS2 Synchronization Findings",
            [
                ("Values successfully pulled", dhis2.get("dhis2_values_successfully_pulled")),
                ("No-data responses", dhis2.get("dhis2_no_data_count")),
                ("Sync errors", dhis2.get("dhis2_error_count")),
            ],
            max_value=max(
                (_coerce_number(dhis2.get("dhis2_values_successfully_pulled")) or 0)
                + (_coerce_number(dhis2.get("dhis2_no_data_count")) or 0)
                + (_coerce_number(dhis2.get("dhis2_error_count")) or 0),
                1,
            ),
        )
        if dhis2.get("last_sync_time"):
            p = document.add_paragraph(f"Last DHIS2 sync time: {dhis2['last_sync_time']}")
            p.runs[0].font.size = Pt(8)
            p.runs[0].font.italic = True

    # --- Facility performance table ---
    facility_rows = [
        [
            item.get("facility_name"),
            item.get("dqa_score"),
            item.get("score_category"),
            item.get("exact_count"),
            item.get("critical_count"),
            item.get("open_corrective_actions"),
        ]
        for item in structured.get("facility_score_ranking", [])
    ]
    _add_data_table(
        document,
        "Facility Performance Table",
        ["Facility", "DQA score", "Category", "Exact", "Critical", "Open actions"],
        facility_rows,
    )

    # --- Indicator findings ---
    indicator_rows = [
        [
            item.get("indicator_name"),
            item.get("hmis_code"),
            item.get("exact_match_rate"),
            item.get("major_discrepancy_count"),
            item.get("critical_discrepancy_count"),
            item.get("worst_facilities") if isinstance(item.get("worst_facilities"), str) else ", ".join(item.get("worst_facilities") or []),
        ]
        for item in structured.get("indicator_findings", [])
    ]
    _add_data_table(
        document,
        "Indicator Statistical Findings",
        ["Indicator", "HMIS code", "Exact rate", "Major", "Critical", "Worst facilities"],
        indicator_rows,
    )

    # --- Indicator performance chart ---
    indicator_chart = _chart_indicator_performance(structured)
    if indicator_chart:
        _add_chart_to_document(document, indicator_chart)

    # --- Detailed comparison rows ---
    comparison_rows = [
        [
            row.get("facility_name"),
            row.get("indicator_name"),
            row.get("hmis_code"),
            row.get("register_value"),
            row.get("hmis105_value"),
            row.get("dhis2_value_at_assessment") or row.get("dhis2_value"),
            row.get("severity"),
            row.get("issue_type"),
        ]
        for row in structured.get("comparison_rows", [])
    ]
    _add_data_table(
        document,
        "Detailed Comparison Rows",
        ["Facility", "Indicator", "HMIS", "Register", "HMIS 105", "DHIS2", "Severity", "Issue"],
        comparison_rows,
    )

    # --- Per-facility details (for multi-facility reports with 10+ facilities) ---
    if facility_count >= 10:
        document.add_page_break()
        document.add_heading("Per-Facility Appendix", level=1)
        comparison_rows_all = structured.get("comparison_rows", [])
        # Group by facility
        facilities_map: dict[str, list[dict]] = {}
        for row in comparison_rows_all:
            fname = row.get("facility_name", "Unknown")
            if fname not in facilities_map:
                facilities_map[fname] = []
            facilities_map[fname].append(row)

        for fname, frows in sorted(facilities_map.items()):
            document.add_heading(f"Facility: {fname}", level=2)
            f_table_rows = [
                [
                    row.get("indicator_name"),
                    row.get("hmis_code"),
                    row.get("register_value"),
                    row.get("hmis105_value"),
                    row.get("dhis2_value_at_assessment") or row.get("dhis2_value"),
                    row.get("severity"),
                    row.get("issue_type"),
                ]
                for row in frows
            ]
            _add_data_table(
                document,
                f"Comparison Details – {fname}",
                ["Indicator", "HMIS", "Register", "HMIS 105", "DHIS2", "Severity", "Issue"],
                f_table_rows,
            )


# ====================================================================
# Narrative rendering
# ====================================================================

def _add_report_narrative(document: Document, content: str) -> None:
    document.add_heading("Legacy Narrative Appendix", level=1)
    for raw_line in content.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            h = document.add_heading(line[2:].strip(), level=1)
            for run in h.runs:
                run.font.color.rgb = UCMB_PRIMARY_BLUE
        elif line.startswith("## "):
            h = document.add_heading(line[3:].strip(), level=2)
            for run in h.runs:
                run.font.color.rgb = UCMB_PRIMARY_BLUE
        elif line.startswith("### "):
            h = document.add_heading(line[4:].strip(), level=3)
            for run in h.runs:
                run.font.color.rgb = UCMB_PRIMARY_BLUE
        elif line.startswith(("- ", "* ")):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            p = document.add_paragraph(line)
            _add_paragraph_spacing(p, before_pt=0, after_pt=6)


# ====================================================================
# V4 finding-block report renderer
# ====================================================================

def _as_text_list(value) -> str:
    if isinstance(value, list):
        return ", ".join(str(item) for item in value if str(item)) or "Not available"
    return str(value) if value else "Not available"


def _add_action_box(document: Document, finding: dict) -> None:
    table = document.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Required action", finding.get("required_action")),
        ("Owner role", finding.get("owner_role")),
        ("Proposed timeline", finding.get("proposed_timeline")),
        ("Evidence required for closure", finding.get("evidence_required_for_closure")),
    ]
    for row_idx, (label, value) in enumerate(rows):
        cells = table.rows[row_idx].cells
        cells[0].text = label
        cells[0].paragraphs[0].runs[0].font.bold = True
        _set_cell_shading(cells[0], UCMB_LIGHT_BLUE_BG)
        cells[1].text = str(value or "Not available")
    document.add_paragraph()


def _add_dhis2_classification_table(document: Document, structured: dict) -> None:
    classification = (structured.get("dhis2_sync_summary") or {}).get("response_classification") or {}
    rows = [
        ["Value returned", classification.get("VALUE_RETURNED", 0), "Stored DHIS2 value was available for comparison."],
        ["True zero", classification.get("TRUE_ZERO", 0), "DHIS2 returned a stored numeric zero."],
        ["No data returned", classification.get("NO_DATA", 0), "No stored value was returned and this cannot be treated as zero without verification."],
        ["Sync error", classification.get("SYNC_ERROR", 0), "DHIS2/API/configuration issue affected extraction."],
        ["Not applicable", classification.get("NOT_APPLICABLE", 0), "Row was not applicable where explicitly classified."],
        ["Unknown", classification.get("UNKNOWN", 0), "Response status was not available or could not be classified."],
    ]
    _add_data_table(document, "DHIS2 Response Classification", ["Classification", "Count", "Interpretation"], rows, limit=10)


def _add_source_document_review_table(document: Document, structured: dict) -> None:
    if structured.get("source_document_assessment_status") == "NOT_ASSESSED" or not structured.get("source_document_checks"):
        p = document.add_paragraph("Source document quality was not fully measured in this round.")
        p.runs[0].font.italic = True
        _add_data_table(
            document,
            "Next-Round Source Document Requirements",
            ["Required check", "Purpose"],
            [
                ["Register availability", "Confirm that the source register was present for verification."],
                ["Register completeness", "Confirm pages and reporting period entries are complete."],
                ["Register legibility", "Confirm entries can be audited."],
                ["Monthly summary presence", "Confirm monthly totals were prepared from register entries."],
                ["Report sign-off", "Confirm facility review and approval before submission."],
                ["HMIS 105 copy availability", "Confirm the submitted HMIS 105 report copy is available."],
                ["HMIS 108 copy availability where applicable", "Confirm related monthly summary forms are available when relevant."],
            ],
        )
        return
    rows = [
        [
            item.get("source_document_name"),
            "Yes" if item.get("available") else "No",
            "Yes" if item.get("complete") else "No",
            "Yes" if item.get("legible") else "No",
            item.get("missing_pages"),
        ]
        for item in structured.get("source_document_checks") or []
    ]
    _add_data_table(document, "Source Document Checklist", ["Document", "Available", "Complete", "Legible", "Missing pages"], rows)


def _render_v4_executive_snapshot(document: Document, structured: dict, blocks: dict) -> None:
    _add_section_heading(document, "Executive Snapshot")
    snapshot = blocks.get("executive_snapshot") or {}
    for key in ("headline", "primary_finding", "management_implication"):
        _add_narrative_paragraphs(document, snapshot.get(key))
    urgent = snapshot.get("urgent_actions") or []
    if urgent:
        _add_data_table(document, "Urgent Management Actions", ["Action"], [[item] for item in urgent], limit=10)


def _render_critical_chase_list(document: Document, structured: dict, blocks: dict) -> None:
    rows = structured.get("critical_chase_list") or []
    if not rows:
        return
    _add_section_heading(document, "Critical Chase List")
    _add_narrative_paragraphs(document, blocks.get("critical_chase_list_intro"))
    _add_data_table(
        document,
        "Critical Death/High-Risk Rows for Reconciliation",
        ["Facility", "Administrative area", "Indicator", "HMIS code", "Register", "HMIS 105", "DHIS2", "Gap", "Pattern", "Owner role", "Proposed target date", "Evidence required"],
        [
            [
                row.get("facility"),
                row.get("administrative_area"),
                row.get("indicator"),
                row.get("hmis_code"),
                row.get("register_value"),
                row.get("hmis105_value"),
                row.get("dhis2_value"),
                row.get("gap"),
                row.get("pattern"),
                row.get("owner_role"),
                row.get("proposed_target_date"),
                row.get("evidence_required_for_closure"),
            ]
            for row in rows
        ],
        limit=20,
    )


def _render_scope_and_method(document: Document, structured: dict, blocks: dict) -> None:
    _add_section_heading(document, "Scope and Method")
    scope = blocks.get("scope_and_method") or {}
    _add_narrative_paragraphs(document, scope.get("scope_summary"))
    _add_narrative_paragraphs(document, scope.get("method_summary"))
    _add_narrative_paragraphs(document, scope.get("denominator_note"))
    _add_narrative_paragraphs(document, scope.get("severity_note"))
    coverage = structured.get("coverage") or {}
    _add_key_value_table(
        document,
        "Coverage Summary",
        [
            ("Facilities selected", coverage.get("total_facilities_selected")),
            ("Facilities assessed", coverage.get("facilities_assessed")),
            ("Facilities pending", coverage.get("facilities_pending")),
            ("Completion percent", coverage.get("percentage_completed")),
            ("Administrative areas covered", ", ".join(coverage.get("administrative_areas_covered") or coverage.get("districts_covered") or [])),
            ("Actual district-level metadata", "Not available in the structured input"),
            ("Facility types", ", ".join(coverage.get("facility_types") or [])),
        ],
    )


def _render_v4_findings(document: Document, structured: dict, blocks: dict, facility_count: int) -> None:
    _add_section_heading(document, "Main Findings")
    for finding in blocks.get("findings") or []:
        title = f"Finding {finding.get('finding_number')}: {finding.get('finding_title')}"
        document.add_heading(title, level=2)
        _add_narrative_paragraphs(document, f"Evidence: {finding.get('evidence')}")
        _add_narrative_paragraphs(document, f"Interpretation: {finding.get('interpretation')}")
        _add_key_value_table(
            document,
            "Affected Scope",
            [
                ("Affected facilities", _as_text_list(finding.get("affected_facilities"))),
                ("Affected indicators", _as_text_list(finding.get("affected_indicators"))),
                ("Affected administrative areas", _as_text_list(finding.get("affected_administrative_areas"))),
                ("Risk level", finding.get("risk_level")),
            ],
        )
        _add_narrative_paragraphs(document, f"Implication: {finding.get('implication')}")
        _add_action_box(document, finding)

        title_lower = str(finding.get("finding_title", "")).lower()
        category_lower = str(finding.get("finding_category", "")).lower()
        if "overall" in title_lower or "quality" in category_lower:
            score = structured.get("dqa_score") or {}
            _add_score_card_table(
                document,
                "Severity Distribution Supporting This Finding",
                [
                    ("Exact matches", score.get("exact_count")),
                    ("Minor discrepancies", score.get("minor_count")),
                    ("Moderate discrepancies", score.get("moderate_count")),
                    ("Major discrepancies", score.get("major_count")),
                    ("Critical discrepancies", score.get("critical_count")),
                    ("Missing values", score.get("missing_count")),
                ],
            )
            chart = _chart_severity_distribution(structured)
            if chart:
                _add_chart_to_document(document, chart)
        if "summarization" in title_lower or "pathway" in category_lower:
            chart = _chart_discrepancy_types(structured)
            if chart:
                _add_chart_to_document(document, chart)
        if "indicator" in title_lower:
            chart = _chart_indicator_performance(structured)
            if chart:
                _add_chart_to_document(document, chart)
        if "facility" in title_lower:
            chart = _chart_facility_scores(structured)
            if chart and facility_count > 1:
                _add_chart_to_document(document, chart)
        if "dhis2" in title_lower or "synchronization" in category_lower:
            _add_dhis2_classification_table(document, structured)
        if "source document" in title_lower:
            _add_source_document_review_table(document, structured)


def _render_v4_facility_performance(document: Document, structured: dict, blocks: dict, facility_count: int) -> None:
    _add_section_heading(document, "Facility Performance")
    _add_narrative_paragraphs(document, (blocks.get("facility_performance_summary") or {}).get("summary"))
    if facility_count > 1:
        chart = _chart_facility_scores(structured)
        if chart:
            _add_chart_to_document(document, chart)
    rows = [
        [
            item.get("facility_name"),
            item.get("administrative_area") or item.get("district"),
            item.get("dqa_score"),
            item.get("score_category"),
            item.get("exact_count"),
            item.get("major_count"),
            item.get("critical_count"),
            item.get("missing_count"),
        ]
        for item in structured.get("facility_score_ranking") or []
    ]
    _add_data_table(document, "Facility Action Matrix", ["Facility", "Administrative area", "Score", "Category", "Exact", "Major", "Critical", "Missing"], rows)


def _render_v4_indicator_performance(document: Document, structured: dict, blocks: dict) -> None:
    _add_section_heading(document, "Indicator Performance")
    _add_narrative_paragraphs(document, (blocks.get("indicator_performance_summary") or {}).get("summary"))
    chart = _chart_indicator_performance(structured)
    if chart:
        _add_chart_to_document(document, chart)
    rows = [
        [
            item.get("hmis_code"),
            item.get("indicator_name"),
            item.get("total_rows") or item.get("n"),
            item.get("exact_match_rate"),
            item.get("major_discrepancy_count"),
            item.get("critical_discrepancy_count"),
            "Review definition/coding and reconcile repeated discrepancies.",
        ]
        for item in structured.get("indicator_findings") or []
    ]
    _add_data_table(document, "Indicator Action Matrix", ["HMIS code", "Indicator", "N", "Exact rate", "Major", "Critical", "Recommended action"], rows)


def _render_v4_corrective_action_plan(document: Document, structured: dict, blocks: dict) -> None:
    _add_section_heading(document, "Corrective Action Plan")
    plan = blocks.get("corrective_action_plan") or {}
    _add_narrative_paragraphs(document, plan.get("summary"))
    rows = [
        [
            action.get("action_id"),
            action.get("linked_finding"),
            action.get("facility_or_scope"),
            action.get("indicator_or_area"),
            action.get("action"),
            action.get("owner_role"),
            action.get("proposed_target_date"),
            action.get("evidence_required_for_closure"),
            action.get("status"),
        ]
        for action in plan.get("actions") or structured.get("ai_corrective_actions") or []
    ]
    _add_data_table(
        document,
        "30/60/90-Day Corrective Action Tracker",
        ["Action ID", "Linked finding", "Facility/scope", "Indicator/area", "Action", "Owner role", "Proposed target date", "Evidence required", "Status"],
        rows,
        limit=50,
    )


def _render_v4_tail_sections(document: Document, structured: dict, blocks: dict) -> None:
    _add_section_heading(document, "DHIS2 No-Data Review")
    review = blocks.get("dhis2_no_data_review") or {}
    _add_narrative_paragraphs(document, review.get("summary"))
    _add_narrative_paragraphs(document, review.get("interpretation"))
    _add_narrative_paragraphs(document, review.get("required_platform_fix"))
    _add_dhis2_classification_table(document, structured)

    _add_section_heading(document, "Source Document Review")
    review = blocks.get("source_document_review") or {}
    _add_narrative_paragraphs(document, review.get("summary"))
    _add_narrative_paragraphs(document, review.get("interpretation"))
    _add_narrative_paragraphs(document, review.get("next_round_requirement"))
    _add_source_document_review_table(document, structured)

    _render_v4_corrective_action_plan(document, structured, blocks)

    _add_section_heading(document, "Limitations")
    for item in blocks.get("limitations") or ["Limitations were not available in the structured data."]:
        document.add_paragraph(str(item), style="List Bullet")

    _add_section_heading(document, "Next Round Improvements")
    for item in blocks.get("next_round_improvements") or ["Next-round improvements were not available in the structured data."]:
        document.add_paragraph(str(item), style="List Bullet")

    _add_section_heading(document, "Appendix Notes")
    _add_narrative_paragraphs(document, "Full detailed comparison rows are available in the companion XLSX export.")
    if structured.get("include_comments"):
        _add_narrative_paragraphs(document, "Detailed field comments are not dumped in the main report. Sanitized comments are available in the companion XLSX export.")
    else:
        _add_narrative_paragraphs(document, "Field comments were excluded from this generated management report.")

    _add_section_heading(document, "Conclusion")
    _add_narrative_paragraphs(document, blocks.get("conclusion"))


def _render_v4_finding_block_report(document: Document, report: Report, facility_count: int) -> None:
    structured = report.structured_input_json or {}
    blocks = structured.get("finding_blocks") or {}
    _render_v4_executive_snapshot(document, structured, blocks)
    _render_critical_chase_list(document, structured, blocks)
    _render_scope_and_method(document, structured, blocks)
    _render_v4_findings(document, structured, blocks, facility_count)
    _render_v4_facility_performance(document, structured, blocks, facility_count)
    _render_v4_indicator_performance(document, structured, blocks)
    _render_v4_tail_sections(document, structured, blocks)


# ====================================================================
# Export log helper
# ====================================================================

def _log_export(
    db: Session,
    *,
    report: Report,
    current_user: User,
    export_type: ExportType,
    file_name: str,
    status_value: ExportStatus,
    error_message: str | None = None,
) -> None:
    db.add(
        ExportLog(
            report_id=report.id,
            exported_by_user_id=current_user.id,
            export_type=export_type,
            file_name=file_name,
            status=status_value,
            error_message=error_message,
            exported_at=datetime.now(UTC),
        )
    )
    if status_value == ExportStatus.SUCCESS:
        report.status = ReportStatus.EXPORTED
        report.exported_by_user_id = current_user.id
        report.exported_at = datetime.now(UTC)
    db.flush()


# ====================================================================
# Public export entry points
# ====================================================================

def export_report_docx(db: Session, report: Report, current_user: User) -> tuple[str, bytes, str]:
    _ensure_export_allowed(report, ExportType.DOCX)
    structured = report.structured_input_json or {}
    facility_count = _count_facilities(structured)
    has_blended_sections = bool(structured.get("narrative_sections"))

    document = Document()
    _add_deepseek_header(document)
    _add_footer(document)

    # ── Cover page ──
    _add_cover_page(document, report, facility_count)

    # ── Metric boxes (headline KPIs) ──
    _add_metric_boxes(document, report)

    if structured.get("finding_blocks"):
        _render_v4_finding_block_report(document, report, facility_count)
    elif has_blended_sections:
        # New path: AI returned structured per-section narrative. Walk each section
        # and interleave prose with tables and charts so the document reads as one
        # coherent narrative rather than dashboard + essay.
        _build_key_findings_section(document, structured, facility_count)
        _render_blended_report(document, report, facility_count)
    else:
        # Legacy path: old reports (or AI failed) use the previous layout where the
        # whole dashboard sits before a single narrative dump at the end.
        _add_statistical_dashboard(document, report, facility_count)
        document.add_page_break()
        _build_key_findings_section(document, structured, facility_count)
        _build_field_notes_section(document, structured, facility_count)
        _build_recommendations_section(document, structured, facility_count)
        document.add_page_break()
        _add_report_narrative(document, _current_report_content(report))

    output = BytesIO()
    document.save(output)
    file_name = _safe_filename(report, ExportType.DOCX)
    _log_export(db, report=report, current_user=current_user, export_type=ExportType.DOCX, file_name=file_name, status_value=ExportStatus.SUCCESS)
    return file_name, output.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def export_report_xlsx(db: Session, report: Report, current_user: User) -> tuple[str, bytes, str]:
    _ensure_export_allowed(report, ExportType.XLSX)
    structured = report.structured_input_json or {}
    blocks = structured.get("finding_blocks") or {}
    score = structured.get("dqa_score") or {}
    summary = structured.get("summary") or {}
    dhis2_summary = structured.get("dhis2_sync_summary") or {}
    response_classification = dhis2_summary.get("response_classification") or {}

    workbook = Workbook()
    dashboard = workbook.active
    dashboard.title = "Executive Dashboard"

    def header(sheet, columns: list[str]) -> None:
        sheet.append(columns)
        for cell in sheet[1]:
            cell.font = Font(bold=True)

    header(dashboard, ["metric", "value", "interpretation"])
    dashboard.append(["report_title", report.title, "Generated management report title."])
    dashboard.append(["report_type", report.report_type.value, "Report type selected by the user."])
    dashboard.append(["status", report.status.value, "Report lifecycle status."])
    dashboard.append(["overall_score_or_exact_rate", score.get("score_percent") or summary.get("exact_match_rate"), "Higher values indicate stronger source alignment."])
    dashboard.append(["critical_count", score.get("critical_count") or summary.get("critical_discrepancy_count"), "Critical rows require immediate reconciliation."])
    dashboard.append(["dhis2_no_data_count", response_classification.get("NO_DATA") or dhis2_summary.get("dhis2_no_data_count"), "No-data is not a true zero without verification."])
    dashboard.append(["source_document_status", structured.get("source_document_assessment_status", "UNKNOWN"), "NOT_ASSESSED means do not interpret completeness as 0%."])

    facility_sheet = workbook.create_sheet("Facility Ranking")
    header(facility_sheet, ["facility", "administrative_area", "dqa_score", "category", "exact_count", "minor_count", "moderate_count", "major_count", "critical_count", "flagged_count", "main_issue", "action_priority"])
    for item in structured.get("facility_score_ranking") or []:
        flagged = (item.get("moderate_count") or 0) + (item.get("major_count") or 0) + (item.get("critical_count") or 0) + (item.get("missing_count") or 0)
        facility_sheet.append([
            item.get("facility_name"),
            item.get("administrative_area") or item.get("district"),
            item.get("dqa_score"),
            item.get("score_category"),
            item.get("exact_count"),
            item.get("minor_count"),
            item.get("moderate_count"),
            item.get("major_count"),
            item.get("critical_count"),
            flagged,
            "Review major, critical, missing, and no-data rows.",
            "High" if (item.get("critical_count") or 0) else "Medium" if flagged else "Routine",
        ])

    indicator_sheet = workbook.create_sheet("Indicator Ranking")
    header(indicator_sheet, ["hmis_code", "indicator_name", "n", "exact_rate", "major_count", "critical_count", "all_zero_note", "interpretation", "recommended_action"])
    for item in structured.get("indicator_findings") or []:
        indicator_sheet.append([
            item.get("hmis_code"),
            item.get("indicator_name"),
            item.get("total_rows") or item.get("n"),
            item.get("exact_match_rate"),
            item.get("major_discrepancy_count"),
            item.get("critical_discrepancy_count"),
            "Interpret cautiously if all or nearly all values are zero.",
            "Prioritize indicators with low exact rate, major discrepancies, or critical discrepancies.",
            "Review definition/coding and reconcile repeated discrepancies.",
        ])

    chase_sheet = workbook.create_sheet("Critical Chase List")
    header(chase_sheet, ["facility", "administrative_area", "indicator", "hmis_code", "register_value", "hmis105_value", "dhis2_value", "gap", "pattern", "owner_role", "proposed_target_date", "evidence_required_for_closure"])
    for row in structured.get("critical_chase_list") or []:
        chase_sheet.append([
            row.get("facility"),
            row.get("administrative_area"),
            row.get("indicator"),
            row.get("hmis_code"),
            row.get("register_value"),
            row.get("hmis105_value"),
            row.get("dhis2_value"),
            row.get("gap"),
            row.get("pattern"),
            row.get("owner_role"),
            row.get("proposed_target_date"),
            row.get("evidence_required_for_closure"),
        ])

    action_sheet = workbook.create_sheet("Corrective Action Tracker")
    header(action_sheet, ["action_id", "linked_finding", "facility_or_scope", "indicator_or_area", "action", "owner_role", "proposed_target_date", "evidence_required_for_closure", "status"])
    for action in (structured.get("ai_corrective_actions") or (blocks.get("corrective_action_plan") or {}).get("actions") or []):
        action_sheet.append([
            action.get("action_id"),
            action.get("linked_finding"),
            action.get("facility_or_scope"),
            action.get("indicator_or_area"),
            action.get("action"),
            action.get("owner_role"),
            action.get("proposed_target_date"),
            action.get("evidence_required_for_closure"),
            action.get("status"),
        ])

    dhis2_sheet = workbook.create_sheet("DHIS2 Sync Audit")
    header(dhis2_sheet, ["facility", "indicator", "hmis_code", "dhis2_value", "dhis2_response_status", "last_sync_time", "sync_error", "interpretation"])
    metadata_by_indicator = {
        item.get("indicator_id"): item for item in structured.get("dhis2_extraction_metadata") or []
    }
    for row in structured.get("comparison_rows") or []:
        metadata = metadata_by_indicator.get(row.get("indicator_id")) or {}
        status_value = row.get("dhis2_response_status") or metadata.get("dhis2_response_status")
        dhis2_sheet.append([
            row.get("facility_name"),
            row.get("indicator_name"),
            row.get("hmis_code"),
            row.get("dhis2_value_at_assessment"),
            status_value,
            metadata.get("dhis2_extracted_at"),
            metadata.get("dhis2_error_message") or row.get("dhis2_error_message"),
            "No-data requires verification and is not true zero." if status_value == "NO_DATA" else "Use with comparison severity and issue type.",
        ])

    source_sheet = workbook.create_sheet("Source Document Checklist")
    header(source_sheet, ["facility", "source_document", "available", "complete", "legible", "missing_pages", "comment", "assessment_status"])
    if structured.get("source_document_checks"):
        for item in structured.get("source_document_checks") or []:
            source_sheet.append([
                item.get("facility_name") or (structured.get("facility") or {}).get("facility_name"),
                item.get("source_document_name"),
                item.get("available"),
                item.get("complete"),
                item.get("legible"),
                item.get("missing_pages"),
                sanitize_comment(item.get("comment")),
                "ASSESSED",
            ])
    else:
        source_sheet.append(["", "", "", "", "", "", "Source document quality was not fully measured in this round.", "NOT_ASSESSED"])

    submitted_sheet = workbook.create_sheet("Submitted Data")
    header(submitted_sheet, ["facility", "administrative_area", "indicator", "hmis_code", "register_value", "hmis105_value", "dhis2_value", "dhis2_response_status", "severity", "issue_type", "comparison_notes"])
    for row in structured.get("comparison_rows") or []:
        submitted_sheet.append([
            row.get("facility_name"),
            row.get("administrative_area") or row.get("district"),
            row.get("indicator_name"),
            row.get("hmis_code"),
            row.get("register_value"),
            row.get("hmis105_value"),
            row.get("dhis2_value_at_assessment"),
            row.get("dhis2_response_status"),
            row.get("severity"),
            row.get("issue_type"),
            row.get("comparison_notes"),
        ])

    comments_sheet = workbook.create_sheet("Field Comments")
    header(comments_sheet, ["facility", "indicator", "comment_type", "sanitized_comment"])
    if structured.get("include_comments"):
        for item in structured.get("general_facility_comments") or []:
            comments_sheet.append([item.get("facility_name"), "", "general", sanitize_comment(item.get("comment"))])
        for item in structured.get("manager_comments") or []:
            comments_sheet.append([item.get("facility_name"), "", "manager", sanitize_comment(item.get("comment"))])
        for row in structured.get("comparison_rows") or []:
            if row.get("assessor_comment"):
                comments_sheet.append([row.get("facility_name"), row.get("indicator_name"), "assessor", sanitize_comment(row.get("assessor_comment"))])
    else:
        comments_sheet.append(["", "", "excluded", "Field comments were excluded from this management report by request/default."])

    dictionary_sheet = workbook.create_sheet("Data Dictionary")
    header(dictionary_sheet, ["field_or_value", "meaning"])
    dictionary_rows = [
        ["administrative_area", "Facility geography as available in the structured input. Do not treat as true district-level metadata unless verified."],
        ["EXACT", "Zero difference across comparable sources."],
        ["MINOR", "Difference within configured tolerance."],
        ["MODERATE", "Above tolerance but not major."],
        ["MAJOR", "Substantial discrepancy requiring action."],
        ["CRITICAL", "Death/high-risk indicator discrepancy or critical row requiring reconciliation."],
        ["MISSING", "One or more required source values missing or unavailable."],
        ["VALUE_RETURNED", "DHIS2 returned a stored value."],
        ["TRUE_ZERO", "DHIS2 returned a stored numeric zero."],
        ["NO_DATA", "DHIS2 returned no stored value; verify before interpreting."],
        ["SYNC_ERROR", "DHIS2 extraction/configuration/API problem."],
        ["source_document_assessment_status", "NOT_ASSESSED means no checklist records existed, not a 0% score."],
        ["proposed_target_date", "AI/template proposed timeline; not an official deadline unless approved by management."],
    ]
    for row in dictionary_rows:
        dictionary_sheet.append(row)

    for sheet in workbook.worksheets:
        for column_cells in sheet.columns:
            length = max(len(str(cell.value or "")) for cell in column_cells)
            sheet.column_dimensions[column_cells[0].column_letter].width = min(max(length + 2, 12), 45)

    output = BytesIO()
    workbook.save(output)
    file_name = _safe_filename(report, ExportType.XLSX)
    _log_export(db, report=report, current_user=current_user, export_type=ExportType.XLSX, file_name=file_name, status_value=ExportStatus.SUCCESS)
    return file_name, output.getvalue(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


def export_report_pdf(db: Session, report: Report, current_user: User) -> tuple[str, bytes, str]:
    _ensure_export_allowed(report, ExportType.PDF)
    file_name = _safe_filename(report, ExportType.PDF)
    if canvas is None or A4 is None:
        _log_export(
            db,
            report=report,
            current_user=current_user,
            export_type=ExportType.PDF,
            file_name=file_name,
            status_value=ExportStatus.FAILED,
            error_message="PDF export dependency is not installed. Install reportlab to enable PDF export.",
        )
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="PDF export dependency is not installed. Install reportlab to enable PDF export.",
        )

    output = BytesIO()
    pdf = canvas.Canvas(output, pagesize=A4)
    width, height = A4
    y = height - 50
    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(50, y, report.title)
    y -= 30
    pdf.setFont("Helvetica", 10)
    for line in _current_report_content(report).splitlines():
        if y < 50:
            pdf.showPage()
            pdf.setFont("Helvetica", 10)
            y = height - 50
        pdf.drawString(50, y, line[:120])
        y -= 14
    pdf.save()
    _log_export(db, report=report, current_user=current_user, export_type=ExportType.PDF, file_name=file_name, status_value=ExportStatus.SUCCESS)
    return file_name, output.getvalue(), "application/pdf"
